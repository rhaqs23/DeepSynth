from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

def export_docx(report: str):
    doc = Document()
    
    # Add title
    title = doc.add_heading("Deep Research Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add report content - split by markdown headings and paragraphs
    current_heading_level = None
    for line in report.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        
        # Detect markdown headings
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('- '):
            # Bullet point
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped[0].isdigit() and '. ' in stripped:
            # Numbered list
            try:
                num, text = stripped.split('. ', 1)
                doc.add_paragraph(text, style='List Number')
            except:
                doc.add_paragraph(stripped)
        else:
            # Regular paragraph
            doc.add_paragraph(stripped)

    buffer = BytesIO()
    doc.save(buffer)

    buffer.seek(0)
    return buffer.getvalue()